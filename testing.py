from docx import Document

doc = Document("data/job_description.docx")

print("Paragraphs:")
print("-" * 50)

for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(i, para.text)

print("\nTables:", len(doc.tables))